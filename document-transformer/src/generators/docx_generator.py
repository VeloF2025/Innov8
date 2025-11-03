"""
Word Document Generator for Document Transformation
Converts business documents to professional Microsoft Word documents
"""

import os
from pathlib import Path
from typing import Optional, Dict, List, Any, Tuple, Union
from dataclasses import dataclass
import logging

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

from ..parser import ParsedDocument, ContentSection, FinancialData, TeamMember
from ..branding import BrandProfile, ColorPalette, Typography

@dataclass
class DocumentOptions:
    """Word document generation options"""
    page_margins: Dict[str, float] = None
    font_size: int = 11
    line_spacing: float = 1.15
    include_toc: bool = False
    include_headers: bool = True
    include_page_numbers: bool = True
    include_cover_page: bool = True

class WordGenerator:
    """Generates professional Word documents from parsed content"""

    def __init__(self, template_engine, output_dir: Path):
        self.template_engine = template_engine
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Setup logging
        self.logger = logging.getLogger(__name__)

    def generate_document(self, document: ParsedDocument, brand_profile: BrandProfile,
                         template_config=None,
                         doc_options: Optional[DocumentOptions] = None,
                         output_filename: Optional[str] = None) -> str:
        """Generate Word document"""

        # Use default options if not provided
        if doc_options is None:
            doc_options = self._get_default_document_options(document)

        # Generate filename if not provided
        if output_filename is None:
            output_filename = self._generate_filename(document)

        # Create document
        doc = Document()

        # Set document properties
        self._set_document_properties(doc, document, brand_profile)

        # Apply styling
        self._apply_document_styling(doc, brand_profile, doc_options)

        # Add content
        self._add_document_content(doc, document, brand_profile, doc_options)

        # Save document
        output_path = self.output_dir / output_filename
        doc.save(str(output_path))

        print(f"✅ Generated Word document: {output_path}")
        return str(output_path)

    def _set_document_properties(self, doc: Document, document: ParsedDocument, brand_profile: BrandProfile):
        """Set document properties"""

        core_props = doc.core_properties

        core_props.title = document.metadata.title or f"{document.metadata.company_name} Document"
        core_props.subject = f"{document.metadata.document_type.replace('_', ' ').title()}"
        core_props.author = document.metadata.author or brand_profile.company_name
        core_props.company = brand_profile.company_name
        core_props.keywords = f"{document.metadata.document_type}, {document.metadata.industry}, business"

        # Add creation date
        from datetime import datetime
        core_props.created = datetime.now()

    def _apply_document_styling(self, doc: Document, brand_profile: BrandProfile, doc_options: DocumentOptions):
        """Apply styling to the document"""

        # Set page margins
        if doc_options.page_margins:
            sections = doc.sections
            for section in sections:
                section.left_margin = Inches(doc_options.page_margins.get('left', 1.0))
                section.right_margin = Inches(doc_options.page_margins.get('right', 1.0))
                section.top_margin = Inches(doc_options.page_margins.get('top', 1.0))
                section.bottom_margin = Inches(doc_options.page_margins.get('bottom', 1.0))

        # Apply brand colors
        self._apply_brand_colors_to_styles(doc, brand_profile.color_palette)

        # Apply brand fonts
        self._apply_brand_fonts_to_styles(doc, brand_profile.typography)

        # Create custom styles
        self._create_custom_styles(doc, brand_profile)

    def _apply_brand_colors_to_styles(self, doc: Document, color_palette: ColorPalette):
        """Apply brand colors to document styles"""

        try:
            # Apply primary color to headings
            if color_palette.primary and len(color_palette.primary) > 0:
                primary_color = self._hex_to_rgb(color_palette.primary[0])

                # Modify heading styles
                for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
                    if style_name in doc.styles:
                        style = doc.styles[style_name]
                        if style.font:
                            style.font.color.rgb = primary_color

        except Exception as e:
            self.logger.warning(f"Could not apply brand colors: {e}")

    def _apply_brand_fonts_to_styles(self, doc: Document, typography: Typography):
        """Apply brand fonts to document styles"""

        try:
            # Apply fonts to styles
            font_mapping = {
                'Heading 1': typography.heading_font,
                'Heading 2': typography.heading_font,
                'Heading 3': typography.heading_font,
                'Normal': typography.body_font
            }

            for style_name, font_name in font_mapping.items():
                if style_name in doc.styles:
                    style = doc.styles[style_name]
                    if style.font:
                        style.font.name = font_name

        except Exception as e:
            self.logger.warning(f"Could not apply brand fonts: {e}")

    def _create_custom_styles(self, doc: Document, brand_profile: BrandProfile):
        """Create custom styles for the document"""

        try:
            # Create custom heading style
            styles = doc.styles

            # Company name style
            company_style = styles.add_style('CompanyName', WD_STYLE_TYPE.PARAGRAPH)
            company_font = company_style.font
            company_font.name = brand_profile.typography.heading_font
            company_font.size = Pt(24)
            company_font.bold = True
            if brand_profile.color_palette.primary:
                company_font.color.rgb = self._hex_to_rgb(brand_profile.color_palette.primary[0])

            # Document title style
            title_style = styles.add_style('DocumentTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = brand_profile.typography.heading_font
            title_font.size = Pt(20)
            title_font.bold = True

            # Section heading style
            section_style = styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
            section_font = section_style.font
            section_font.name = brand_profile.typography.heading_font
            section_font.size = Pt(16)
            section_font.bold = True
            if brand_profile.color_palette.primary:
                section_font.color.rgb = self._hex_to_rgb(brand_profile.color_palette.primary[0])

            # Table header style
            table_header_style = styles.add_style('TableHeader', WD_STYLE_TYPE.PARAGRAPH)
            table_header_font = table_header_style.font
            table_header_font.name = brand_profile.typography.body_font
            table_header_font.size = Pt(11)
            table_header_font.bold = True

        except Exception as e:
            self.logger.warning(f"Could not create custom styles: {e}")

    def _add_document_content(self, doc: Document, document: ParsedDocument,
                             brand_profile: BrandProfile, doc_options: DocumentOptions):
        """Add content to the document"""

        # Add cover page if requested
        if doc_options.include_cover_page:
            self._add_cover_page(doc, document, brand_profile)

        # Add table of contents if requested
        if doc_options.include_toc:
            self._add_table_of_contents(doc, document)

        # Add document sections
        self._add_sections(doc, document.sections, brand_profile)

        # Add financial data if available
        if document.financial_data:
            self._add_financial_data(doc, document.financial_data, brand_profile)

        # Add team information if available
        if document.team_members:
            self._add_team_information(doc, document.team_members, brand_profile)

        # Add headers and footers if requested
        if doc_options.include_headers:
            self._add_headers_footers(doc, document, brand_profile)

    def _add_cover_page(self, doc: Document, document: ParsedDocument, brand_profile: BrandProfile):
        """Add cover page"""

        # Add title
        title = document.metadata.title or f"{document.metadata.company_name} Document"
        title_para = doc.add_paragraph(title)
        title_para.style = 'DocumentTitle'
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add company name
        if document.metadata.company_name:
            company_para = doc.add_paragraph(document.metadata.company_name)
            company_para.style = 'CompanyName'
            company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add document type
        doc_type_para = doc.add_paragraph(document.metadata.document_type.replace('_', ' ').title())
        doc_type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add date
        from datetime import datetime
        date_para = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page break
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def _add_table_of_contents(self, doc: Document, document: ParsedDocument):
        """Add table of contents"""

        toc_para = doc.add_paragraph("Table of Contents")
        toc_para.style = 'Heading 1'

        # Add section entries
        for i, section in enumerate(document.sections, 1):
            toc_entry = doc.add_paragraph()
            toc_entry.add_run(f"{i}. {section.title}")
            toc_entry.paragraph_format.left_indent = Inches(0.25)

        # Add page break
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def _add_sections(self, doc: Document, sections: List[ContentSection], brand_profile: BrandProfile):
        """Add document sections"""

        for section in sections:
            # Add section heading
            heading_para = doc.add_paragraph(section.title)
            heading_para.style = 'SectionHeading'

            # Add section content
            if section.content:
                # Split content into paragraphs
                content_paragraphs = section.content.split('\n')
                for para_text in content_paragraphs:
                    if para_text.strip():
                        para = doc.add_paragraph(para_text.strip())
                        para.style = 'Normal'

            # Add spacing between sections
            doc.add_paragraph()  # Empty paragraph for spacing

    def _add_financial_data(self, doc: Document, financial_data: List[FinancialData], brand_profile: BrandProfile):
        """Add financial data tables"""

        # Add section heading
        finance_heading = doc.add_paragraph("Financial Information")
        finance_heading.style = 'SectionHeading'

        for financial in financial_data:
            if financial.title:
                table_title = doc.add_paragraph(financial.title)
                table_title.style = 'Heading 3'

            if financial.table_data and len(financial.table_data) > 0:
                # Create table
                table = doc.add_table(rows=len(financial.table_data), cols=len(financial.table_data[0]))
                table.style = 'Medium Grid 1 Accent 1'

                # Add data to table
                for row_idx, row_data in enumerate(financial.table_data):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < len(financial.table_data[0]):  # Ensure we don't exceed column count
                            cell = table.cell(row_idx, col_idx)
                            cell.text = str(cell_data)

                            # Style header row
                            if row_idx == 0:
                                cell.paragraphs[0].style = 'TableHeader'

                # Add some spacing after table
                doc.add_paragraph()

    def _add_team_information(self, doc: Document, team_members: List[TeamMember], brand_profile: BrandProfile):
        """Add team member information"""

        # Add section heading
        team_heading = doc.add_paragraph("Leadership Team")
        team_heading.style = 'SectionHeading'

        # Add team members
        for member in team_members:
            # Add member name
            name_para = doc.add_paragraph(member.name)
            name_para.style = 'Heading 4'

            # Add member title
            if member.title:
                title_para = doc.add_paragraph(member.title)
                title_para.style = 'Strong'

            # Add member bio
            if member.bio:
                bio_para = doc.add_paragraph(member.bio)
                bio_para.style = 'Normal'

            # Add experience if available
            if member.experience:
                exp_para = doc.add_paragraph(f"Experience: {member.experience}")
                exp_para.style = 'Normal'

            # Add spacing between team members
            doc.add_paragraph()

    def _add_headers_footers(self, doc: Document, document: ParsedDocument, brand_profile: BrandProfile):
        """Add headers and footers to the document"""

        try:
            section = doc.sections[0]

            # Add header
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = f"{brand_profile.company_name} - {document.metadata.document_type.replace('_', ' ').title()}"
            header_para.style = 'Header'

            # Add footer
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "Page "
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        except Exception as e:
            self.logger.warning(f"Could not add headers/footers: {e}")

    def _get_default_document_options(self, document: ParsedDocument) -> DocumentOptions:
        """Get default document options based on document type"""

        doc_type = document.metadata.document_type

        base_options = DocumentOptions()

        if doc_type in ['business_plan', 'market_research']:
            base_options.include_toc = True
            base_options.include_cover_page = True
            base_options.include_headers = True
        elif doc_type == 'investor_teaser':
            base_options.include_cover_page = True
            base_options.include_headers = False
        elif doc_type == 'pitch_deck':
            base_options.page_margins = {'left': 0.75, 'right': 0.75, 'top': 1.0, 'bottom': 1.0}
        else:
            base_options.include_cover_page = True

        return base_options

    def _hex_to_rgb(self, hex_color: str) -> RGBColor:
        """Convert hex color to RGBColor object"""

        hex_color = hex_color.lstrip('#')
        if len(hex_color) == 6:
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16)
            b = int(hex_color[4:6], 16)
            return RGBColor(r, g, b)
        return RGBColor(0, 0, 0)  # Black fallback

    def _generate_filename(self, document: ParsedDocument) -> str:
        """Generate appropriate filename for Word document"""

        company = document.metadata.company_name.lower().replace(' ', '_')
        project = document.metadata.project.lower().replace(' ', '_') if document.metadata.project else ""
        doc_type = document.metadata.document_type.lower().replace('_', '-')
        title = document.metadata.title.lower().replace(' ', '_')[:30] if document.metadata.title else ""

        parts = [part for part in [company, project, doc_type, title] if part]
        filename = "_".join(parts) + ".docx"

        # Sanitize filename
        import re
        filename = re.sub(r'[^\w\-_\.]', '', filename)

        return filename

    def generate_batch_documents(self, documents: List[Tuple[ParsedDocument, BrandProfile]],
                               doc_options: Optional[DocumentOptions] = None) -> List[str]:
        """Generate Word documents for multiple documents"""

        output_paths = []

        for document, brand_profile in documents:
            try:
                output_path = self.generate_document(
                    document, brand_profile, None, doc_options
                )
                output_paths.append(output_path)
            except Exception as e:
                print(f"❌ Error generating Word document for {document.metadata.title}: {e}")
                continue

        return output_paths

    def create_summary_document(self, documents: List[ParsedDocument],
                               brand_profile: BrandProfile,
                               output_filename: str = "document_summary.docx") -> str:
        """Create a summary document containing information about all documents"""

        doc = Document()

        # Set document properties
        core_props = doc.core_properties
        core_props.title = f"{brand_profile.company_name} Document Summary"
        core_props.author = brand_profile.company_name

        # Apply styling
        self._apply_document_styling(doc, brand_profile, DocumentOptions())

        # Add title
        title_para = doc.add_paragraph(f"{brand_profile.company_name} Document Summary")
        title_para.style = 'DocumentTitle'
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add date
        from datetime import datetime
        date_para = doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page break
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # Add summary table
        summary_heading = doc.add_paragraph("Document Summary")
        summary_heading.style = 'SectionHeading'

        # Create summary table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Medium Grid 1 Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Add headers
        headers = ["Document Title", "Type", "Sections", "Date Created"]
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].style = 'TableHeader'

        # Add document information
        for document in documents:
            row_cells = table.add_row().cells
            row_cells[0].text = document.metadata.title or "Untitled"
            row_cells[1].text = document.metadata.document_type.replace('_', ' ').title()
            row_cells[2].text = str(len(document.sections))
            row_cells[3].text = document.metadata.last_modified or "Unknown"

        # Add document details
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        details_heading = doc.add_paragraph("Document Details")
        details_heading.style = 'SectionHeading'

        for document in documents:
            # Document title
            doc_title = doc.add_paragraph(document.metadata.title or "Untitled Document")
            doc_title.style = 'Heading 3'

            # Document metadata
            metadata_para = doc.add_paragraph()
            metadata_para.add_run(f"Type: ").bold = True
            metadata_para.add_run(f"{document.metadata.document_type.replace('_', ' ').title()}\n")
            metadata_para.add_run(f"Company: ").bold = True
            metadata_para.add_run(f"{document.metadata.company_name}\n")
            metadata_para.add_run(f"Industry: ").bold = True
            metadata_para.add_run(f"{document.metadata.industry.title()}\n")
            metadata_para.add_run(f"Sections: ").bold = True
            metadata_para.add_run(f"{len(document.sections)}")

            # Brief description
            if document.sections:
                desc_para = doc.add_paragraph()
                desc_para.add_run("Description: ").bold = True
                desc_para.add_run(document.sections[0].content[:200] + "..." if len(document.sections[0].content) > 200 else document.sections[0].content)

            # Add spacing
            doc.add_paragraph()

        # Save document
        output_path = self.output_dir / output_filename
        doc.save(str(output_path))

        print(f"✅ Generated summary document: {output_path}")
        return str(output_path)